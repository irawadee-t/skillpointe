"""
Resume upload + LLM parse for applicant onboarding.

Pure-ish helpers: read PDF/DOCX bytes, extract text, feed the existing
`applicant_extractor` pipeline. The router persists results.
"""
from __future__ import annotations

import io
import re
import sys
from pathlib import Path
from typing import Any

# packages/extraction lives outside apps/api. Add it to sys.path IF present.
# Walk only existing parents (avoids IndexError when the deploy layout is
# flatter than the repo — e.g. Railway Root Directory = apps/api).
for _parent in Path(__file__).resolve().parents:
    _pkg = _parent / "packages"
    if _pkg.is_dir():
        if str(_pkg) not in sys.path:
            sys.path.insert(0, str(_pkg))
        break


MAX_UPLOAD_BYTES = 4 * 1024 * 1024  # 4 MB


def extract_text_from_upload(content: bytes, content_type: str, filename: str) -> str:
    """Return raw text from an uploaded resume. Empty string on unsupported type."""
    ct = (content_type or "").lower()
    fn = (filename or "").lower()
    if "pdf" in ct or fn.endswith(".pdf"):
        return _extract_pdf(content)
    if "word" in ct or "officedocument" in ct or fn.endswith(".docx"):
        return _extract_docx(content)
    if "text" in ct or fn.endswith(".txt"):
        return content.decode("utf-8", errors="replace")
    return ""


def _extract_pdf(content: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(content))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    return _clean(("\n".join(parts)))


def _extract_docx(content: bytes) -> str:
    import docx
    d = docx.Document(io.BytesIO(content))
    parts = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    # Also read text out of tables (many resumes use them for layout).
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return _clean("\n".join(parts))


def _clean(text: str) -> str:
    if not text:
        return ""
    # Normalize whitespace, drop repeated blank lines, cap absurd length.
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()[:60_000]


# ---------------------------------------------------------------------------
# Extract structured signals via the existing LLM extractor + regex fallbacks.
# ---------------------------------------------------------------------------

def parse_resume_signals(text: str, *, openai_api_key: str | None, model: str) -> dict[str, Any]:
    """Return a partial applicant profile suggested by the resume.

    Keys align with the applicants table columns:
      first_name, last_name, phone, email, city, state,
      program_name_raw, experience_raw, career_goals_raw, bio_raw,
      certifications (list of names), skills (list of names).
    """
    signals: dict[str, Any] = {}
    if not text.strip():
        return signals

    # Deterministic parses — fast, cheap, don't need an API key.
    signals.update(_regex_contact(text))

    if openai_api_key:
        from openai import OpenAI
        client = OpenAI(api_key=openai_api_key)
        try:
            llm_out = _call_llm(client, model, text)
            signals.update(llm_out)
        except Exception:
            # Fall through with just the regex-parsed contact fields.
            pass

    return signals


_EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
_PHONE_RE = re.compile(r"(?:\+?1[\s.-]?)?\(?\b\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}\b")
_STATE_RE = re.compile(r"\b([A-Z]{2})\b(?:\s+\d{5})?")  # matches "GA 30303" or trailing "GA"


def _regex_contact(text: str) -> dict[str, Any]:
    out: dict[str, Any] = {}
    lines = [l.strip() for l in text.splitlines()[:15] if l.strip()]

    # Name = first non-empty line with 2-4 title-cased words and no digits/at/URL markers.
    for line in lines:
        if any(bad in line.lower() for bad in ["@", "http", "resume", "cv"]):
            continue
        parts = line.split()
        if 2 <= len(parts) <= 4 and all(p[:1].isupper() and p.replace("-", "").isalpha() for p in parts if p):
            out["first_name"] = parts[0]
            out["last_name"] = " ".join(parts[1:])
            break

    if m := _EMAIL_RE.search(text):
        out["email"] = m.group(0)
    if m := _PHONE_RE.search(text):
        out["phone"] = m.group(0)
    # City/state — look at first 20 lines for "City, ST" pattern.
    for line in text.splitlines()[:20]:
        m = re.search(r"([A-Z][A-Za-z .-]+),\s*([A-Z]{2})\b", line)
        if m:
            out["city"] = m.group(1).strip()
            out["state"] = m.group(2)
            break
    return out


_LLM_SYSTEM = (
    "You are extracting fields from a trades / skilled-worker resume. "
    "Return strict JSON only. Never invent data — if a field is missing, omit it. "
    "Focus on trades relevance: welding, electrical, HVAC, CDL, mechanic, industrial "
    "maintenance, manufacturing operator, technician, etc."
)

_LLM_USER_TEMPLATE = (
    "Extract from this resume:\n"
    "- program_name_raw: the training program or trade specialty (e.g. \"Industrial Electrician\", \"AWS D1.1 Welding\")\n"
    "- career_goals_raw: 1-2 sentence summary of what work they're looking for\n"
    "- experience_raw: brief plain-text summary of their most recent 1-3 roles\n"
    "- bio_raw: any personal statement or summary text at the top of the resume\n"
    "- certifications: list of certification names (e.g. \"OSHA 30\", \"CDL Class A\", \"EPA 608\", \"AWS D1.1\")\n"
    "- skills: list of skill keywords (e.g. \"MIG welding\", \"blueprint reading\", \"forklift operation\")\n"
    "- years_experience: integer years of hands-on experience if inferable\n\n"
    "Return JSON with only the fields you can find. RESUME TEXT:\n\n{text}"
)


def _call_llm(client, model: str, text: str) -> dict[str, Any]:
    from openai import OpenAI  # noqa: F401 (typing)
    resp = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": _LLM_SYSTEM},
            {"role": "user",   "content": _LLM_USER_TEMPLATE.format(text=text[:15_000])},
        ],
        response_format={"type": "json_object"},
        temperature=0.1,
        max_tokens=1200,
    )
    import json
    raw = resp.choices[0].message.content or "{}"
    data = json.loads(raw)

    # Normalize types.
    out: dict[str, Any] = {}
    for k in ("program_name_raw", "career_goals_raw", "experience_raw", "bio_raw"):
        v = data.get(k)
        if isinstance(v, str) and v.strip():
            out[k] = v.strip()
    for k in ("certifications", "skills"):
        v = data.get(k)
        if isinstance(v, list):
            clean = [str(x).strip() for x in v if isinstance(x, (str, int)) and str(x).strip()]
            if clean:
                out[k] = clean[:20]
    years = data.get("years_experience")
    if isinstance(years, (int, float)) and 0 <= years <= 60:
        out["years_experience"] = int(years)
    return out
